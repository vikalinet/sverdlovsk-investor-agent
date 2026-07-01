"""
Модуль экспорта документов в различные форматы
"""
from pathlib import Path
from typing import Optional, BinaryIO
from datetime import datetime
from loguru import logger

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx не установлен. Экспорт в DOCX недоступен.")

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("reportlab не установлен. Экспорт в PDF недоступен.")


class DocumentExporter:
    """
    Экспорт документов в различные форматы
    
    Поддерживаемые форматы:
    - TXT (по умолчанию)
    - DOCX (требуется python-docx)
    - PDF (требуется reportlab)
    """
    
    def __init__(self, output_dir: Optional[Path] = None):
        if output_dir is None:
            output_dir = Path(__file__).parent.parent / "output" / "documents"
        
        self.output_dir = output_dir
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        logger.info(f"DocumentExporter инициализирован: {output_dir}")
    
    def export(
        self,
        content: str,
        filename: str,
        format: str = "txt",
        metadata: Optional[dict] = None
    ) -> Path:
        """
        Экспорт документа
        
        Args:
            content: Содержимое документа
            filename: Имя файла (без расширения)
            format: Формат (txt, docx, pdf)
            metadata: Метаданные документа
        
        Returns:
            Путь к сохранённому файлу
        """
        format = format.lower()
        
        if format == "txt":
            return self._export_txt(content, filename, metadata)
        elif format == "docx" and DOCX_AVAILABLE:
            return self._export_docx(content, filename, metadata)
        elif format == "pdf" and PDF_AVAILABLE:
            return self._export_pdf(content, filename, metadata)
        else:
            logger.error(f"Формат {format} не поддерживается")
            raise ValueError(f"Формат {format} не поддерживается или зависимость не установлена")
    
    def _export_txt(
        self,
        content: str,
        filename: str,
        metadata: Optional[dict] = None
    ) -> Path:
        """Экспорт в TXT"""
        output_path = self.output_dir / f"{filename}.txt"
        
        # Добавление метаданных в начало
        header = self._create_metadata_header(metadata)
        full_content = header + "\n" + content if header else content
        
        output_path.write_text(full_content, encoding='utf-8')
        logger.info(f"Документ экспортирован: {output_path}")
        return output_path
    
    def _export_docx(
        self,
        content: str,
        filename: str,
        metadata: Optional[dict] = None
    ) -> Path:
        """Экспорт в DOCX"""
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx не установлен")
        
        output_path = self.output_dir / f"{filename}.docx"
        
        doc = Document()
        
        # Заголовок
        if metadata and 'title' in metadata:
            title = doc.add_heading(metadata['title'], 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Метаданные
        if metadata:
            doc.add_paragraph()
            for key, value in metadata.items():
                if key != 'title':
                    p = doc.add_paragraph()
                    p.add_run(f"{key}: ").bold = True
                    p.add_run(str(value))
        
        # Основное содержание
        doc.add_page_break()
        
        # Парсинг текста по абзацам
        for line in content.split('\n'):
            if line.strip():
                doc.add_paragraph(line)
        
        # Сохранение
        doc.save(str(output_path))
        logger.info(f"Документ экспортирован: {output_path}")
        return output_path
    
    def _export_pdf(
        self,
        content: str,
        filename: str,
        metadata: Optional[dict] = None
    ) -> Path:
        """Экспорт в PDF"""
        if not PDF_AVAILABLE:
            raise RuntimeError("reportlab не установлен")
        
        output_path = self.output_dir / f"{filename}.pdf"
        
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=A4,
            rightMargin=0.75*inch,
            leftMargin=0.75*inch,
            topMargin=0.75*inch,
            bottomMargin=0.75*inch
        )
        
        story = []
        styles = getSampleStyleSheet()
        
        # Заголовок
        if metadata and 'title' in metadata:
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=30,
                alignment=WD_ALIGN_PARAGRAPH.CENTER
            )
            story.append(Paragraph(metadata['title'], title_style))
            story.append(Spacer(1, 0.3*inch))
        
        # Метаданные
        if metadata:
            for key, value in metadata.items():
                if key != 'title':
                    p = Paragraph(f"<b>{key}:</b> {value}", styles['Normal'])
                    story.append(p)
            story.append(Spacer(1, 0.3*inch))
        
        # Основное содержание
        for line in content.split('\n'):
            if line.strip():
                # Обработка форматирования
                if line.startswith('='):
                    story.append(Spacer(1, 0.2*inch))
                else:
                    p = Paragraph(line.replace('&', '&amp;'), styles['Normal'])
                    story.append(p)
            story.append(Spacer(1, 0.1*inch))
        
        # Сохранение
        doc.build(story)
        logger.info(f"Документ экспортирован: {output_path}")
        return output_path
    
    def _create_metadata_header(self, metadata: Optional[dict]) -> str:
        """Создание заголовка с метаданными"""
        if not metadata:
            return ""
        
        lines = [
            "=" * 80,
            f"Документ сгенерирован: {datetime.now().strftime('%d.%m.%Y %H:%M')}",
            "=" * 80,
            ""
        ]
        
        for key, value in metadata.items():
            lines.append(f"{key}: {value}")
        
        lines.append("")
        lines.append("=" * 80)
        lines.append("")
        
        return "\n".join(lines)
    
    def export_multiple(
        self,
        documents: list[dict],
        format: str = "txt"
    ) -> list[Path]:
        """
        Экспорт нескольких документов
        
        Args:
            documents: Список документов [{'content': str, 'filename': str, 'metadata': dict}]
            format: Формат экспорта
        
        Returns:
            Список путей к файлам
        """
        paths = []
        for doc in documents:
            path = self.export(
                content=doc['content'],
                filename=doc['filename'],
                format=format,
                metadata=doc.get('metadata')
            )
            paths.append(path)
        
        return paths
    
    def create_archive(
        self,
        documents: list[Path],
        archive_name: str
    ) -> Path:
        """
        Создание архива с документами
        
        Args:
            documents: Список путей к документам
            archive_name: Имя архива
        
        Returns:
            Путь к архиву
        """
        import zipfile
        
        archive_path = self.output_dir / f"{archive_name}.zip"
        
        with zipfile.ZipFile(archive_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for doc_path in documents:
                if doc_path.exists():
                    zipf.write(doc_path, doc_path.name)
        
        logger.info(f"Архив создан: {archive_path}")
        return archive_path
